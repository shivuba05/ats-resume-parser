"""Extractor module for extracting raw text and embedded links from PDF and DOCX files.

Supports pdfminer.six, pypdf, and python-docx with robust pure-Python fallbacks.
"""

import io
import os
from pathlib import Path
import re
from typing import BinaryIO, Union
import xml.etree.ElementTree as ET
import zipfile
import zlib

try:
    from pdfminer.high_level import extract_text as extract_pdfminer_text
    from pdfminer.layout import LAParams
except ImportError:
    extract_pdfminer_text = None
    LAParams = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None


class UnsupportedFormatError(ValueError):
    """Raised when an unsupported file format is provided."""
    pass


def _extract_pdf_annotations_urls(data: bytes) -> list[str]:
    """Extract embedded hyperlink URLs from PDF annotation objects."""
    urls = []
    # Match /URI (http...) in PDF dictionary objects
    uri_matches = re.findall(rb"/URI\s*\((https?://[^\)]+)\)", data, re.IGNORECASE)
    for u in uri_matches:
        try:
            urls.append(u.decode("utf-8", errors="ignore"))
        except Exception:
            pass
    return list(dict.fromkeys(urls))


def _extract_text_with_pypdf(data: bytes) -> tuple[str, list[str]]:
    """Extract text and links from PDF bytes using pypdf."""
    if pypdf is None:
        return "", []

    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        text_pages = []
        links = []

        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_pages.append(t)

            # Extract links from annotations if present
            if "/Annots" in page:
                try:
                    for annot in page["/Annots"]:
                        obj = annot.get_object()
                        if obj and "/A" in obj and "/URI" in obj["/A"]:
                            uri = str(obj["/A"]["/URI"])
                            if uri.startswith("http"):
                                links.append(uri)
                except Exception:
                    pass

        return "\n".join(text_pages), list(dict.fromkeys(links))
    except Exception:
        return "", []


def _extract_text_from_docx_xml(file_source: Union[str, Path, BinaryIO, bytes]) -> str:
    """Fallback DOCX text extractor using built-in zipfile and ElementTree XML parsing."""
    try:
        if isinstance(file_source, (str, Path)):
            zf = zipfile.ZipFile(str(file_source))
        elif isinstance(file_source, bytes):
            zf = zipfile.ZipFile(io.BytesIO(file_source))
        else:
            zf = zipfile.ZipFile(file_source)

        with zf.open("word/document.xml") as doc_xml:
            tree = ET.parse(doc_xml)
            root = tree.getroot()

            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

            paragraphs = []
            for p in root.iter(f"{{{ns['w']}}}p"):
                texts = [node.text for node in p.iter(f"{{{ns['w']}}}t") if node.text]
                p_text = "".join(texts).strip() if texts else ""
                paragraphs.append(p_text)

            return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX content: {e}") from e


def _extract_text_from_pdf_fallback(data: bytes) -> str:
    """Fallback basic PDF text stream extractor using built-in zlib and regex."""
    try:
        extracted_chunks = []

        stream_matches = re.finditer(rb"stream[\r\n]+(.*?)[\r\n]+endstream", data, re.DOTALL)
        for m in stream_matches:
            stream_data = m.group(1)
            decompressed = None
            try:
                decompressed = zlib.decompress(stream_data)
            except Exception:
                try:
                    decompressed = zlib.decompress(stream_data, -15)
                except Exception:
                    decompressed = stream_data

            if decompressed:
                text_matches = re.findall(rb"\((.*?)\)\s*(?:Tj|['\"])", decompressed)
                for t in text_matches:
                    clean_str = t.decode("latin1", errors="ignore").replace("\\(", "(").replace("\\)", ")").replace("\\\\", "\\")
                    extracted_chunks.append(clean_str)

                tj_blocks = re.findall(rb"\[(.*?)\]\s*TJ", decompressed)
                for block in tj_blocks:
                    block_texts = re.findall(rb"\((.*?)\)", block)
                    if block_texts:
                        combined = "".join(bt.decode("latin1", errors="ignore").replace("\\(", "(").replace("\\)", ")").replace("\\\\", "\\") for bt in block_texts)
                        extracted_chunks.append(combined)

        embedded_urls = _extract_pdf_annotations_urls(data)
        if embedded_urls:
            extracted_chunks.append("\n" + "\n".join(embedded_urls))

        if extracted_chunks:
            return "\n".join(extracted_chunks)

        text_content = data.decode("latin1", errors="ignore")
        clean_lines = [line.strip() for line in text_content.split("\n") if len(line.strip()) > 3]
        return "\n".join(clean_lines)
    except Exception as e:
        raise ValueError(f"Failed fallback PDF extraction: {e}") from e


def extract_text_from_pdf(file_source: Union[str, Path, BinaryIO, bytes]) -> str:
    """Extract raw text from a PDF file path or file-like stream using best available extractor."""
    data = b""
    if isinstance(file_source, (str, Path)):
        file_path = str(file_source)
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        with open(file_path, "rb") as f:
            data = f.read()
    elif isinstance(file_source, bytes):
        data = file_source
    elif hasattr(file_source, "read"):
        pos = file_source.tell() if hasattr(file_source, "tell") else 0
        data = file_source.read()
        if hasattr(file_source, "seek"):
            file_source.seek(pos)

    text = ""
    links = []

    # 1. Primary: pypdf (fast and handles modern font encodings & layout)
    if pypdf is not None:
        pypdf_text, pypdf_links = _extract_text_with_pypdf(data)
        if pypdf_text and len(pypdf_text.strip()) > 20:
            text = pypdf_text
            links = pypdf_links

    # 2. Secondary: pdfminer.six (handles complex continuous multi-column flows)
    if not text and extract_pdfminer_text is not None and LAParams is not None:
        try:
            laparams = LAParams(
                line_margin=0.5,
                word_margin=0.1,
                char_margin=2.0,
                boxes_flow=0.5,
                detect_vertical=False,
                all_texts=True
            )
            pdfminer_text = extract_pdfminer_text(io.BytesIO(data), laparams=laparams) or ""
            if pdfminer_text and len(pdfminer_text.strip()) > 20:
                text = pdfminer_text
        except Exception:
            pass

    # 3. Fallback: custom stream reader
    if not text:
        text = _extract_text_from_pdf_fallback(data)

    # Append embedded links if not already present
    annot_links = _extract_pdf_annotations_urls(data)
    all_links = list(dict.fromkeys(links + annot_links))
    for url in all_links:
        if url not in text:
            text += f"\n{url}"

    return text


def extract_text_from_docx(file_source: Union[str, Path, BinaryIO, bytes]) -> str:
    """Extract raw text from a DOCX file path or file-like stream."""
    if isinstance(file_source, (str, Path)):
        file_path = str(file_source)
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

    if docx is not None:
        try:
            if isinstance(file_source, (str, Path)):
                doc = docx.Document(str(file_source))
            elif isinstance(file_source, bytes):
                doc = docx.Document(io.BytesIO(file_source))
            else:
                doc = docx.Document(file_source)

            lines = []
            for paragraph in doc.paragraphs:
                lines.append(paragraph.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        lines.append(" | ".join(row_cells))

            return "\n".join(lines)
        except Exception:
            return _extract_text_from_docx_xml(file_source)
    else:
        return _extract_text_from_docx_xml(file_source)


def extract_text(file_source: Union[str, Path, BinaryIO, bytes], file_name: str = "") -> str:
    """Extract raw text from either PDF or DOCX file source.

    Args:
        file_source: File path, file-like object, or raw bytes.
        file_name: Optional explicit filename to deduce format.

    Returns:
        Extracted raw text.

    Raises:
        UnsupportedFormatError: If the file format is not .pdf or .docx.
        FileNotFoundError: If the specified file does not exist.
        ValueError: If extraction fails.
    """
    extension = ""

    if isinstance(file_source, (str, Path)):
        file_path = Path(file_source)
        if not file_path.exists():
            raise FileNotFoundError(f"Resume file not found at: {file_source}")
        extension = file_path.suffix.lower()
    elif hasattr(file_source, "name") and file_source.name:
        extension = Path(file_source.name).suffix.lower()
    elif file_name:
        extension = Path(file_name).suffix.lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file_source)
    elif extension == ".docx":
        return extract_text_from_docx(file_source)
    elif extension == ".doc":
        raise UnsupportedFormatError(
            "Legacy '.doc' format is not directly supported. Please convert the file to '.docx' or '.pdf'."
        )
    elif extension:
        raise UnsupportedFormatError(
            f"Unsupported file format '{extension}'. Only PDF (.pdf) and Word (.docx) files are supported."
        )
    else:
        header = b""
        if isinstance(file_source, bytes):
            header = file_source[:10]
        elif hasattr(file_source, "read") and hasattr(file_source, "seek"):
            current_pos = file_source.tell()
            header = file_source.read(10)
            file_source.seek(current_pos)

        if header.startswith(b"%PDF"):
            return extract_text_from_pdf(file_source)
        elif header.startswith(b"PK\x03\x04"):
            return extract_text_from_docx(file_source)
        else:
            raise UnsupportedFormatError(
                "Unable to detect file format. Please provide a valid .pdf or .docx file."
            )
